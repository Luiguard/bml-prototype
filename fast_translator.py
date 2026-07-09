import os
import sys
import subprocess
from pdf2docx import Converter
from docx import Document

in_pdf = sys.argv[1]
out_pdf = sys.argv[2]
temp_docx = in_pdf + ".docx"

# Dictionary für die schnelle Direktübersetzung ohne LLM
vocab = {
    "Süss": "Sladké", "süss": "sladké", "Süß": "Sladké", "süß": "sladké",
    "Fruchtig": "Ovocné", "fruchtig": "ovocné",
    "Blumig": "Kvetinové", "blumig": "kvetinové",
    "Aromatisch": "Aromatické", "aromatisch": "aromatické",
    "Zitrisch": "Citrusové", "zitrisch": "citrusové",
    "Cremig": "Krémové", "cremig": "krémové",
    "Orientalisch": "Orientálne", "orientalisch": "orientálne",
    "Frisch": "Svieže", "frisch": "svieže",
    "Aquatisch": "Vodné", "aquatisch": "vodné",
    "Würzig": "Korenisté", "würzig": "korenisté",
    "Ledrig": "Kožené", "ledrig": "kožené",
    "Holzig": "Drevité", "holzig": "drevité",
    "Erdig": "Zemité", "erdig": "zemité",
    "Synthetisch": "Syntetické", "synthetisch": "syntetické",
    "Grün": "Zelené", "grün": "zelené",
    "Harzig": "Živicové", "harzig": "živicové",
    "Rauchig": "Dymové", "rauchig": "dymové",
    "Pudrig": "Púdrové", "pudrig": "púdrové",
    "Gourmandig": "Gurmánske", "gourmandig": "gurmánske",
    "Duschgel": "Sprchový gél",
    "Bodylotion": "Telové mlieko",
    "Herren": "Pánske",
    "Damen": "Dámske",
    "Unisex": "Unisex",
    "Eigenkreationen": "Vlastné kreácie",
    "inspiriert von": "inšpirované"
}

def translate_fast(text):
    for de, sk in vocab.items():
        text = text.replace(de, sk)
    return text

print("Extrahiere Layout & Bilder...")
try:
    cv = Converter(in_pdf)
    cv.convert(temp_docx)
    cv.close()
except Exception as e:
    print(f"PDF2DOCX Error: {e}")
    sys.exit(1)

print("Übersetze Textblöcke...")
try:
    doc = Document(temp_docx)
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = translate_fast(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para.text = translate_fast(para.text)
    
    translated_docx = in_pdf + "_translated.docx"
    doc.save(translated_docx)
except Exception as e:
    print(f"DOCX Translation Error: {e}")
    sys.exit(1)

print("Rekonstruiere PDF mit originalen Bildern...")
try:
    outdir = os.path.dirname(out_pdf) or '.'
    subprocess.run(['soffice', '--headless', '--convert-to', 'pdf', translated_docx, '--outdir', outdir], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    gen_pdf = os.path.join(outdir, os.path.splitext(os.path.basename(translated_docx))[0] + ".pdf")
    os.rename(gen_pdf, out_pdf)
except Exception as e:
    print(f"LibreOffice Error: {e}")
    sys.exit(1)
finally:
    try:
        os.remove(temp_docx)
        os.remove(translated_docx)
    except:
        pass

print("Fertig!")
